import pandas as pd
from docx import Document  # word相关

def getHTML():
    file_path = r"C:\Users\420\Downloads\RDS_GNSS.xlsx"
    df = pd.read_excel(file_path, usecols=["latitude", "longtitude"])
    df_sorted = df.loc[179:]
    lines = df_sorted.shape[0]  # 2750
    document = Document()
    for i in range(lines):
        # 以下为绘制线，参考https://blog.csdn.net/fallwind_of_july/article/details/100145473
        # print("new BMap.Point(" + str(df_sorted.iloc[i,1])+ "," + str(df_sorted.iloc[i,0]) + ")," )
        # document.add_paragraph("new BMap.Point(" + str(df_sorted.iloc[i,1])+ "," + str(df_sorted.iloc[i,0]) + "),")

        print("{\"lng\":"+ str(df_sorted.iloc[i,1]) + ",\"lat\":" + str(df_sorted.iloc[i,0]) + "},")
        document.add_paragraph("{\"lng\":"+ str(df_sorted.iloc[i,1]) + ",\"lat\":" + str(df_sorted.iloc[i,0]) + "},")

    document.save(r'C:\Users\420\Downloads\out.docx')



if __name__ == '__main__':
    getHTML()

